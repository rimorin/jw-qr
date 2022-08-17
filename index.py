import io
import requests
import qrcode
import lxml
import cchardet
import re
import pyshorteners
import os
from random import choice
from urllib.parse import urlparse
from bs4 import BeautifulSoup, SoupStrainer
from PIL import Image, ImageOps, ImageDraw, ImageFont
from flask import (
    Flask,
    request,
    render_template,
    send_file,
    abort,
    json,
    send_from_directory,
)
from docx import Document
from docx.shared import Inches, Mm
import logging
import time
import segno
from qrcode.image.styledpil import StyledPilImage
from qrcode.image.styles.moduledrawers import (
    VerticalBarsDrawer,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s %(threadName)s : %(message)s",
)


TINYURL_TIMEOUT_SECONDS = 5
IMAGE_TAG = "og:image"
TITLE_TAG = "og:title"
IMAGE_OFFSET = 45
ARTICLE_IMAGE_SIZE = (160, 160)
ARTICLE_QR_SIZE = (180, 180)
ARTICLE_IMG_QR_SIZE = (120, 120)
TITLE_Y_POS = 8
TITLE_IGNORE_KEYS = ["awake", "watchtower", "videos"]
EXPECTED_DOMAIN = "www.jw.org"
TITLE_LENGTH_THRESHOLD = 60
DEFAULT_FONT = "assets/fonts/NotoSans-Bold.ttf"
URL_LENGTH_THRESHOLD = 180

MEDIAITEM_URL_KEYS = ["mediaitems"]
FINDER_MEDIAITEM_URL_KEYS = ["finder", "VIDEO"]
DEFAULT_MEDIA_IMAGE = "https://assetsnffrgf-a.akamaihd.net/assets/m/802013129/univ/art/802013129_univ_sqr_md.jpg"
DEFAULT_MEDIA_BANNER = "https://assetsnffrgf-a.akamaihd.net/assets/m/802013129/univ/art/802013129_univ_wsr_md.jpg"

DOC_ROWS = 9
DOC_COLUMNS = 4

PROXY_API_URL = "https://ephemeral-proxies.p.rapidapi.com/v1/proxy"
PROXY_API_KEY = os.environ.get("PROXY_API_KEY", "")
PROXY_TIMEOUT = 5
RAPID_API_HEADER = {
    "X-RapidAPI-Host": "ephemeral-proxies.p.rapidapi.com",
    "X-RapidAPI-Key": PROXY_API_KEY,
}

SCRAPER_API_KEY = os.environ.get("SCRAPER_API_KEY", "")
SCRAPER_API_ENDPT = os.environ.get("SCRAPER_API_ENDPT", "http://api.scrape.do/")

CHROME_BROWSER_AGENTS = [
    "Mozilla/5.0 (iPod; CPU iPhone OS 15_5 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) CriOS/103.0.5060.63 Mobile/15E148 Safari/604.1",
    "Mozilla/5.0 (iPad; CPU OS 15_5 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) CriOS/103.0.5060.63 Mobile/15E148 Safari/604.1",
    "Mozilla/5.0 (iPhone; CPU iPhone OS 15_5 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) CriOS/103.0.5060.63 Mobile/15E148 Safari/604.1",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 12_4) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/74.0.3729.169 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3626.121 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/74.0.3729.157 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.140 Safari/537.36 Edge/18.17763",
]

DEFAULT_HEADER = {
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "none",
    "Sec-Fetch-User": "?1",
    "Cache-Control": "max-age=0",
    "Referer": "http://www.google.com/",
}

DEFAULT_QR_TEMPLATE_DESIGN_CODE = 1

app = Flask(__name__)
logger = app.logger
shortener = pyshorteners.Shortener(timeout=TINYURL_TIMEOUT_SECONDS)


def gen_doc(img):
    document = Document()
    section = document.sections[0]
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(10.4)
    section.right_margin = Mm(10.4)
    section.top_margin = Mm(10.4)
    section.bottom_margin = Mm(10.4)
    table = document.add_table(rows=DOC_ROWS, cols=DOC_COLUMNS)
    for row in range(len(table.rows)):
        for col in range(len(table.columns)):
            paragraph = table.cell(row, col).paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(img, height=Inches(1))
    word_file = io.BytesIO()
    document.save(word_file)
    word_file.seek(0)
    return word_file


def qr_processor(
    article_link="", article_title="", article_design=DEFAULT_QR_TEMPLATE_DESIGN_CODE
):
    return eval(f"gen_qr_{article_design}")(article_link, article_title)


def gen_qr(article_link="", article_title="", design=None):
    if not article_link:
        return

    result = urlparse(article_link)
    domain_link = result.netloc.casefold()
    if domain_link != EXPECTED_DOMAIN:
        logger.warn("Invalid domain.")
        abort(404, "Please enter a link from {EXPECTED_DOMAIN}.")
    links = {}
    try:
        links = scrape_article(article_link=article_link)
    except Exception as er:
        logger.error("Error during scraping.", exc_info=True)
        abort(
            404,
            description=f"Opps!! Something is wrong somewhere. Please try another link.",
        )
    
    logoName = "assets/images/logo-1.png"
    if design == 1:
        logoName = "assets/images/logo-2.png"
    elif design == 2:
        logoName = "assets/images/logo-3.png"

    logo = Image.open(logoName)
    left_image = logo.resize(ARTICLE_IMAGE_SIZE, Image.ANTIALIAS)
    left_image = add_margin(
        left_image, top=0, bottom=15, left=15, right=15, color=(250, 250, 250)
    )
    right_image = get_qr_image(
        article_link=article_link, with_logo=False
    )
    article_size = left_image.size
    complete_qr_image = Image.new(
        "RGB", (2 * article_size[0], article_size[1] + IMAGE_OFFSET), (250, 250, 250)
    )
    complete_qr_image.paste(left_image, (0, IMAGE_OFFSET))
    complete_qr_image.paste(right_image, (article_size[0], IMAGE_OFFSET - 10))

    qr_title = article_title if article_title else links.get("title", "")

    draw_title(
        image=complete_qr_image,
        width=article_size[0],
        title=qr_title,
        lang=links.get("lang", "en"),
    )
    complete_qr_image = draw_border(image=complete_qr_image)
    qr_file = io.BytesIO()
    # complete_qr_image.save("test.jpg", "JPEG", quality=95)
    complete_qr_image.save(qr_file, "JPEG", quality=95)
    qr_file.seek(0)
    return qr_file


def gen_qr_1(article_link="", article_title=""):
    return gen_qr(article_link=article_link, article_title=article_title)

def gen_qr_2(article_link="", article_title=""):
    return gen_qr(
        article_link=article_link,
        article_title=article_title,
        design=1,
    )

def gen_qr_3(article_link="", article_title=""):
    return gen_qr(
        article_link=article_link,
        article_title=article_title,
        design=2,
    )


# def gen_qr_2(article_link="", article_title=""):
#     if not article_link:
#         return

#     result = urlparse(article_link)
#     domain_link = result.netloc.casefold()
#     if domain_link != EXPECTED_DOMAIN:
#         logger.warn("Invalid domain.")
#         abort(404, "Please enter a link from {EXPECTED_DOMAIN}.")
#     links = {}
#     try:
#         links = scrape_article(article_link=article_link)
#     except Exception as er:
#         logger.error("Error during scraping.", exc_info=True)
#         abort(
#             404,
#             description=f"Opps!! Something is wrong somewhere. Please try another link.",
#         )
#     left_image = get_article_image(
#         image_url=links.get("image", ""), article_url=article_link
#     )
#     logo = Image.open("assets/images/siteLogo-jworg.png")
#     right_image = logo.resize(ARTICLE_IMAGE_SIZE, Image.ANTIALIAS)
#     article_size = left_image.size
#     complete_qr_image = Image.new(
#         "RGB", (2 * article_size[0], article_size[1] + IMAGE_OFFSET), (250, 250, 250)
#     )
#     complete_qr_image.paste(left_image, (0, IMAGE_OFFSET))
#     complete_qr_image.paste(right_image, (article_size[0] + 10, IMAGE_OFFSET))

#     qr_title = article_title if article_title else links.get("title", "")

#     draw_title(
#         image=complete_qr_image,
#         width=article_size[0],
#         title=qr_title,
#         lang=links.get("lang", "en"),
#     )
#     complete_qr_image = draw_border(image=complete_qr_image)
#     qr_file = io.BytesIO()
#     complete_qr_image.save(qr_file, "JPEG", quality=95)
#     qr_file.seek(0)
#     return qr_file


# def gen_qr_4(article_link="", article_title=""):
#     if not article_link:
#         return

#     result = urlparse(article_link)
#     domain_link = result.netloc.casefold()
#     if domain_link != EXPECTED_DOMAIN:
#         logger.warn("Invalid domain.")
#         abort(404, "Please enter a link from {EXPECTED_DOMAIN}.")
#     links = {}
#     try:
#         links = scrape_article(article_link=article_link)
#     except Exception as er:
#         logger.error("Error during scraping.", exc_info=True)
#         abort(
#             404,
#             description=f"Opps!! Something is wrong somewhere. Please try another link.",
#         )
#     response = get_link_data(link=links.get("banner", ""))
#     bg_file = io.BytesIO(response.content)

#     banner_image = Image.open(bg_file)
#     banner_image = banner_image.resize(
#         ((ARTICLE_IMAGE_SIZE[0] * 2) - IMAGE_OFFSET, ARTICLE_IMAGE_SIZE[1]),
#         Image.ANTIALIAS,
#     )
#     cropped_banner_image = banner_image.crop(
#         (
#             banner_image.size[0] - ARTICLE_IMG_QR_SIZE[0],
#             banner_image.size[1] - ARTICLE_IMG_QR_SIZE[1],
#             banner_image.size[0],
#             banner_image.size[1],
#         )
#     )
#     cropped_image = io.BytesIO()
#     cropped_banner_image.save(cropped_image, format="PNG")
#     cropped_qr_image = segno.make(
#         prepare_link(article_link=article_link, force_shorten=True), error="h"
#     )
#     cropped_qr_file = io.BytesIO()
#     cropped_qr_image.to_artistic(
#         background=cropped_image, target=cropped_qr_file, kind="jpeg", border=0
#     )
#     new_image_file = Image.open(cropped_qr_file)
#     new_image_file = new_image_file.resize(ARTICLE_IMG_QR_SIZE, Image.ANTIALIAS)
#     image_position = (
#         banner_image.size[0] - new_image_file.size[0],
#         banner_image.size[1] - new_image_file.size[1],
#     )
#     banner_image.paste(new_image_file, image_position)
#     left_image = prepare_logo(basewidth=50, border=False)
#     image_position = (0, banner_image.size[1] - left_image.size[1])
#     banner_image.paste(left_image, image_position)

#     qr_title = article_title if article_title else links.get("title", "")
#     draw_title(
#         image=banner_image,
#         width=ARTICLE_IMAGE_SIZE[0] - 25,
#         title=qr_title,
#         lang=links.get("lang", "en"),
#         with_outline=True,
#         title_pos=(0, 3),
#     )
#     banner_image = add_margin(
#         banner_image, top=5, bottom=5, left=5, right=5, color=(250, 250, 250)
#     )
#     banner_image = draw_border(image=banner_image)
#     qr_file = io.BytesIO()
#     # banner_image.save("test.jpg", "JPEG", quality=95)
#     banner_image.save(qr_file, "JPEG", quality=95)
#     qr_file.seek(0)
#     return qr_file


def get_proxy():
    try:
        response = requests.request(
            "GET",
            PROXY_API_URL,
            headers=RAPID_API_HEADER,
        ).json()
        proxy_address = (
            f"http://{response['proxy']['host']}:{response['proxy']['port']}"
        )
        return {"http": proxy_address, "https": proxy_address}
    except:
        logger.error(
            "Error when getting proxy address. Using scraper service.", exc_info=True
        )
        return


def get_link_data(link, use_api=False):
    scrape_header = DEFAULT_HEADER
    scrape_header["User-Agent"] = choice(CHROME_BROWSER_AGENTS)
    NO_API_KEYS_CONFIGURED = not PROXY_API_KEY or not SCRAPER_API_KEY
    if not use_api or NO_API_KEYS_CONFIGURED:
        return requests.get(link)

    proxy_addresses = get_proxy()
    if not proxy_addresses:
        logger.warn("No proxy IP provided. Using external scraper service.")
        # Use scraper.do if no proxy IP is provided
        return requests.get(
            SCRAPER_API_ENDPT, params={"url": link, "token": SCRAPER_API_KEY}
        )

    scrape_response = None
    try:
        logger.info(f"Scraping link={link} with proxy={proxy_addresses}.")
        scrape_response = requests.get(
            link, headers=scrape_header, proxies=proxy_addresses, timeout=PROXY_TIMEOUT
        )
    except Exception as er:
        logger.error(
            f"Detected a problem with the proxy, {proxy_addresses}. Using external scraper service.",
            exc_info=True,
        )
        # Use scraper.do if there is a problem using proxy IP
        scrape_response = requests.get(
            SCRAPER_API_ENDPT, params={"url": link, "token": SCRAPER_API_KEY}
        )
    return scrape_response


def scrape_article(article_link):
    default_link = get_default_link_data(article_link=article_link)
    if default_link:
        return default_link
    page_response = get_link_data(link=article_link, use_api=True)
    return generate_tags(page_response, article_link)


def get_default_link_data(article_link):

    is_finder_link = FINDER_MEDIAITEM_URL_KEYS[0] in article_link
    if is_finder_link:
        if not all(key in article_link for key in FINDER_MEDIAITEM_URL_KEYS):
            return
    else:
        if not any(key in article_link for key in MEDIAITEM_URL_KEYS):
            return

    return {
        "image": DEFAULT_MEDIA_IMAGE,
        "title": "",
        "lang": "en",
        "banner": DEFAULT_MEDIA_BANNER,
    }


def generate_tags(response_data, page_link):
    if not response_data or not response_data.ok:
        # Use scraper.do if proxy response is not ok
        response_data = requests.get(
            SCRAPER_API_ENDPT, params={"url": page_link, "token": SCRAPER_API_KEY}
        )

    tag_data = extract_tags(response_data.text)

    if not tag_data:
        # Use scraper.do if contents from proxy scraper are invalid
        response_data = requests.get(
            SCRAPER_API_ENDPT, params={"url": page_link, "token": SCRAPER_API_KEY}
        )
        tag_data = extract_tags(response_data.text)

    if not tag_data:
        raise Exception(
            "Unable to scrape data using both proxy IP and scraper service."
        )
    return tag_data


def extract_tags(data):
    soup = BeautifulSoup(
        data, "lxml", parse_only=SoupStrainer(["meta", "link", "figure"])
    )

    image_tag = soup.find("meta", property=IMAGE_TAG)
    title_tag = soup.find("meta", property=TITLE_TAG)
    link_tag = soup.find("link", rel="alternate")
    figure_tag = soup.find("figure")

    something_wrong_with_scraped_content = (
        not image_tag or not title_tag or not link_tag
    )

    if something_wrong_with_scraped_content:
        return

    article_image = image_tag["content"]
    banner_image = (
        figure_tag.span.get("data-img-size-md") if figure_tag else article_image
    )
    return {
        "image": article_image,
        "title": title_tag["content"],
        "lang": link_tag.get("hreflang", "en"),
        "banner": banner_image,
    }


def get_article_image(image_url, article_url=None):
    webpage_image_bytes = None
    if article_url:
        response = get_link_data(link=image_url)
        bg_file = io.BytesIO(response.content)
        qrcode = segno.make(
            prepare_link(article_link=article_url, force_shorten=True), error="h"
        )
        webpage_image_bytes = io.BytesIO()
        qrcode.to_artistic(
            background=bg_file, target=webpage_image_bytes, kind="jpeg", border=0
        )
    else:
        response = get_link_data(link=image_url)
        webpage_image_bytes = io.BytesIO(response.content)
    article_image = Image.open(webpage_image_bytes)
    article_image = article_image.resize(ARTICLE_IMAGE_SIZE, Image.ANTIALIAS)
    article_image = add_margin(
        article_image, top=0, bottom=15, left=15, right=15, color=(250, 250, 250)
    )
    return article_image


def prepare_logo(basewidth=140, border=True):
    logo = Image.open("assets/images/siteLogo-jworg.png")
    wpercent = basewidth / float(logo.size[0])
    hsize = int((float(logo.size[1]) * float(wpercent)))
    logo = logo.resize((basewidth, hsize), Image.ANTIALIAS)
    if border == False:
        return logo
    return draw_border(logo, size=(6, 6, 6, 6), color="white")


def prepare_link(article_link, force_shorten=False):
    if len(article_link) > URL_LENGTH_THRESHOLD or force_shorten:
        article_link = shortener.tinyurl.short(article_link)
    return article_link


def get_qr_image(article_link, with_logo=True, design=None):
    QRcode = qrcode.QRCode(error_correction=qrcode.constants.ERROR_CORRECT_H)
    QRcode.add_data(prepare_link(article_link=article_link))
    QRcode.make()
    QRimg = get_design_qr(QRcode, design=design)
    if with_logo:
        logo = prepare_logo()
        pos = ((QRimg.size[0] - logo.size[0]) // 2, (QRimg.size[1] - logo.size[1]) // 2)
        QRimg.paste(logo, pos)
    QRimg = QRimg.resize(ARTICLE_QR_SIZE, Image.ANTIALIAS)
    return QRimg


def get_design_qr(qr, design=None):
    if not design:
        return qr.make_image(back_color=(250, 250, 250)).convert("RGB")

    img = None
    if design == 1:
        img = qr.make_image(
            back_color=(250, 250, 250),
            image_factory=StyledPilImage,
            module_drawer=VerticalBarsDrawer(),
        ).convert("RGB")

    return img


def draw_border(image, size=(2, 2, 2, 2), color="black"):
    return ImageOps.expand(image, border=size, fill=color)


def get_language(lang):

    font_size = 34
    font_type = DEFAULT_FONT

    if "cmn" in lang:
        font_size = 28
        font_type = "assets/fonts/NotoSansSC-Bold.otf"

    if "ta" in lang:
        font_type = "assets/fonts/NotoSansTamil.ttf"

    if "my" in lang:
        font_size = 28
        font_type = "assets/fonts/NotoSansMyanmar-Bold.ttf"

    if "ja" in lang:
        font_size = 28
        font_type = "assets/fonts/NotoSansJP-Bold.otf"

    return font_type, font_size


def draw_title(
    image, width, title, lang, with_outline=False, title_pos=(0, TITLE_Y_POS)
):
    font_type, font_size = get_language(lang=lang)
    font = ImageFont.truetype(font_type, font_size)
    draw = ImageDraw.Draw(image)
    singleline_text(
        draw,
        process_title(title=title),
        font,
        xy=title_pos,
        wh=(2 * width, 30),
        alignment="center",
        with_outline=with_outline,
    )


def process_title(title):
    if "|" in title:
        partition = title.rpartition("|")
        if any(ext in partition[0].casefold() for ext in TITLE_IGNORE_KEYS):
            return partition[2]
        return partition[0]
    elif "—" in title:
        title = title.rpartition("—")[0]
    elif ":" in title:
        title = title.rpartition(":")[2]

    if len(title) > TITLE_LENGTH_THRESHOLD:
        try:
            first_non_alphanumeric_char = re.search(r"\W\s+", title).start()
            title = title[0 : first_non_alphanumeric_char + 1]
        except Exception as e:
            logger.error(
                "Error detected when processing title. Using default title",
                exc_info=True,
            )

    return title


def add_margin(pil_img, top, right, bottom, left, color):
    width, height = pil_img.size
    new_width = width + right + left
    new_height = height + top + bottom
    result = Image.new(pil_img.mode, (new_width, new_height), color)
    result.paste(pil_img, (left, top))
    return result


def singleline_text(
    drawing, text, font_info, xy, wh, fill="#000", alignment=None, with_outline=None
):
    x, y = xy
    container_width, container_height = wh
    x_offset = 0
    y_offset = 0
    font = font_info
    fontsize = font.size
    font_path = font.path
    shadowcolor = "white"

    text_width = font.getsize(text)[0]
    while text_width > container_width:
        # iterate until the text width is smaller than the assigned width of the text area
        fontsize -= 1
        font = ImageFont.truetype(font_path, fontsize)
        text_width = font.getsize(text)[0]

    # optionally de-increment to be sure it is less than criteria
    fontsize -= 1
    font = ImageFont.truetype(font_path, fontsize)
    updated_font_size = font.getsize(text)
    text_width = updated_font_size[0]
    text_height = updated_font_size[1]
    y_offset = (int(container_height) / 2) - (text_height / 2)
    if alignment == "center":
        x_offset = (int(container_width) / 2) - (text_width / 2)
    elif alignment == "right":
        x_offset = container_width - text_width

    final_tuple = (x + x_offset, y + y_offset)
    if with_outline:
        drawing.text(
            (final_tuple[0] - 1, final_tuple[1] - 1), text, font=font, fill=shadowcolor
        )
        drawing.text(
            (final_tuple[0] + 1, final_tuple[1] - 1), text, font=font, fill=shadowcolor
        )
        drawing.text(
            (final_tuple[0] - 1, final_tuple[1] + 1), text, font=font, fill=shadowcolor
        )
        drawing.text(
            (final_tuple[0] + 1, final_tuple[1] + 1), text, font=font, fill=shadowcolor
        )
    drawing.text(final_tuple, text, font=font, fill=fill)


@app.route("/robots.txt")
@app.route("/sitemap.xml")
@app.route("/favicon_io/favicon-32x32.png")
def static_from_root():
    return send_from_directory(app.static_folder, request.path[1:])


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        data = json.loads(request.data)
        article_link = data.get("article_link", "")
        article_title = data.get("article_title", "")
        article_design = data.get("article_design", DEFAULT_QR_TEMPLATE_DESIGN_CODE)
        logger.info(
            f"Attempting to generate QR document. link={article_link}, title={article_title}, design={article_design}"
        )
        process_start_time = time.time()
        img_file = qr_processor(
            article_link=article_link,
            article_title=article_title,
            article_design=article_design,
        )
        word_doc = gen_doc(img=img_file)
        process_end_time = time.time()
        logger.info(
            f"Generated QR document. link={article_link}, title={article_title}, process_time={int(process_end_time-process_start_time)}s"
        )
        return send_file(
            word_doc,
            download_name="article-doc.docx",
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    return render_template("index.jinja2")


if __name__ == "__main__":
    app.run()
